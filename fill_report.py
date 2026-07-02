#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fill the course design report template for Student Score Management System."""
import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATE = r"D:\xiazaiwenjian\java demand\5.《面向对象程序设计(Java)课程设计》- 课程设计报告模板.docx"
OUTPUT   = r"D:\xiazaiwenjian\java demand\StudentScoreSystem_Report_Filled.docx"

doc = Document(TEMPLATE)
paragraphs = doc.paragraphs

def find_para_containing(keyword):
    for i, p in enumerate(paragraphs):
        if keyword in p.text:
            return i, p
    return None, None

def clear_para(index):
    p = paragraphs[index]
    for run in p.runs:
        run.text = ""

# ============================================================
# Fill cover page info
# ============================================================
print("Filling cover page...")

idx, p = find_para_containing("院")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("院（系）：  计算机科学与技术学院")

idx, p = find_para_containing("学号")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("学    号：  2024XXXXXXXX")

idx, p = find_para_containing("姓")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("姓    名：  ______________")

idx, p = find_para_containing("专业班级")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("专业班级：  计科XX-X班")

idx, p = find_para_containing("指导教师")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("指导教师：  ______________")

idx, p = find_para_containing("2026")
if idx is not None and "月" in p.text:
    clear_para(idx)
    paragraphs[idx].add_run("2026 年 6 月 30 日")

# ============================================================
# Part 1: Project Overview
# ============================================================
print("Filling Part 1: Project Overview...")

idx, p = find_para_containing("简单介绍")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "本课程设计实现了一个「学生成绩管理系统」，采用 Java Swing 图形界面与面向对象编程思想，"
        "具备学生信息的增删改查、按总分排序、模糊搜索、正态分布测试数据生成、TXT/Excel 报表导出"
        "以及数据持久化等完整功能。系统界面简洁直观，代码结构清晰，符合 MVC 分层设计思想。"
    )

idx, p = find_para_containing("git地址")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("参考Git地址：https://github.com/Why123-max/why")

idx, p = find_para_containing("个人仓库地址")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("个人仓库地址：https://github.com/Why123-max/why  （独立完成）")

idx, p = find_para_containing("团队仓库地址")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("团队仓库地址：不涉及（独立完成）")

idx, p = find_para_containing("博客地址")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run("个人博客地址：无")

# ============================================================
# Part 2: Architecture Diagram
# ============================================================
print("Filling Part 2: Architecture...")

idx, p = find_para_containing("功能架构图")
if idx is not None:
    clear_para(idx)
    text = (
        "System Architecture (all modules are independently completed):\n\n"
        "+-------------------------------------------------------+\n"
        "|     Student Score Management System (Java Swing)      |\n"
        "+-------------------------------------------------------+\n"
        "|  [LoginGUI]  [StudentGUI]        [ScoreManager]      |\n"
        "|  login       add/delete/update    business logic      |\n"
        "|  verify      search/sort/export   data persistence   |\n"
        "|              [Student.java]       file I/O            |\n"
        "|              entity class         student_data.txt    |\n"
        "|  [JTable + DefaultTableModel]     [BufferedReader/    |\n"
        "|   table display                   BufferedWriter]     |\n"
        "+-------------------------------------------------------+\n"
        "|  [Statistics & Report] [Test Data Generator]          |\n"
        "|  class averages        Random.nextGaussian()          |\n"
        "|  TXT report export     100K records / one click       |\n"
        "+-------------------------------------------------------+"
    )
    paragraphs[idx].add_run(text)

# ============================================================
# Part 3: Technical Introduction
# ============================================================
print("Filling Part 3: Technical Introduction...")

idx, p = find_para_containing("完成的技术与功能")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "This project uses pure Java (no frameworks). Technologies and features:\n\n"
        "1. Java Swing GUI\n"
        "   - JFrame, JPanel layouts (BorderLayout/GridLayout/FlowLayout)\n"
        "   - JTable + DefaultTableModel for data display\n"
        "   - JTextField/JPasswordField for input\n"
        "   - JButton + ActionListener (anonymous inner class)\n"
        "   - JOptionPane for dialogs and confirmations\n\n"
        "2. Object-Oriented Programming (OOP)\n"
        "   - Student entity class (private fields + public getters/setters)\n"
        "   - Constructor overloading (no-arg + parameterized)\n"
        "   - MVC: Model(Student) -> Controller(ScoreManager) -> View(GUI)\n\n"
        "3. Java I/O for Data Persistence\n"
        "   - BufferedReader/FileReader to load data\n"
        "   - BufferedWriter/FileWriter to save data\n"
        "   - CSV format with comma separation (String.split/join)\n\n"
        "4. Collections Framework & Algorithms\n"
        "   - ArrayList<Student> for dynamic storage\n"
        "   - Collections.sort() + Comparator for sorting\n"
        "   - Random.nextGaussian() for normal distribution\n\n"
        "5. Key Features\n"
        "   - 100K test records with normal distribution (mean=80, stdev=10)\n"
        "   - Auto-increment unique student ID (2026 prefix)\n"
        "   - Fuzzy search by name (String.contains)\n"
        "   - Real-time class statistics (count, averages of each subject)"
    )

idx, p = find_para_containing("特色")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "Highlights and Challenges:\n"
        "1. Unique ID mechanism: traverse existing IDs, find max + 1\n"
        "2. Normal distribution: score = 80 + 10 * nextGaussian(), clamped to [0,100]\n"
        "3. Auto-persist: saveDataToFile() called after every add/update/delete\n"
        "4. Fuzzy search: String.contains() for partial name matching\n"
        "5. Instant UI feedback: table auto-refresh after each operation"
    )

idx, p = find_para_containing("Git提交记录")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "Git commit history (1 commit):\n"
        "  commit 46f8309 - Initial commit: student score management system\n"
        "  12 files: .gitignore, Guess.java, LoginGUI.java, Main.java, README.txt,\n"
        "  ScoreManager.java, Student.java, StudentGUI.java, StudentScoreSystem.iml,\n"
        "  guess.txt, student_data.txt, report file\n"
        "  Remote: https://github.com/Why123-max/why"
    )

# ============================================================
# Part 4: Detailed Function Descriptions
# ============================================================
print("Filling Part 4: Detailed Descriptions...")

idx, p = find_para_containing("自己完成的功能")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "This project was completed independently. Below is a detailed description "
        "of each module's implementation, key code, and design rationale."
    )

# 4.1 Login
idx, p = find_para_containing("*登录功能")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "1. Login Module (LoginGUI.java)\n\n"
        "[Description]\n"
        "The login window appears at startup. Users must enter correct credentials "
        "to access the main system. Default account: admin / 123456.\n\n"
        "[Technical Details]\n"
        "- Extends JFrame for independent window\n"
        "- GridLayout(3,2) for labels and input fields\n"
        "- JPasswordField for masked password input\n"
        "- Anonymous ActionListener for login/exit buttons\n"
        "- On success: dispose() login window, new StudentGUI() opens main window\n"
        "- On failure: JOptionPane error message\n\n"
        "[Key Code]\n"
        "loginBtn.addActionListener(new ActionListener() {\n"
        "    public void actionPerformed(ActionEvent e) {\n"
        "        if (username.equals(\"admin\") && password.equals(\"123456\")) {\n"
        "            JOptionPane.showMessageDialog(null, \"Login success!\");\n"
        "            dispose();\n"
        "            new StudentGUI().setVisible(true);\n"
        "        } else {\n"
        "            JOptionPane.showMessageDialog(null, \"Wrong credentials!\");\n"
        "        }\n"
        "    }\n"
        "});"
    )

# 4.2 Core business
idx, p = find_para_containing("*核心业务功能")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "2. Core Business Module (ScoreManager.java)\n\n"
        "[Architecture]\n"
        "ScoreManager is the business logic core - all CRUD operations, statistics, "
        "file I/O, and test data generation. GUI only handles display and events. "
        "Follows MVC pattern.\n\n"
        "[Core Data Structure]\n"
        "private ArrayList<Student> studentList;  // stores all students\n"
        "private int idCounter;                  // ensures unique IDs\n\n"
        "[CRUD Implementation]\n"
        "1. Add: auto-generate ID -> add to ArrayList -> auto-save\n"
        "2. Delete: iterate ArrayList, match ID -> remove -> auto-save\n"
        "3. Update: find index by ID -> set() -> auto-save\n"
        "4. Find by ID: iterate, equals match\n"
        "5. Find by name: iterate, contains() fuzzy match, return list\n\n"
        "[Challenge: Unique ID]\n"
        "On each add, traverse all existing IDs, find max sequence number, add 1.\n"
        "Format: 2026 + 5-digit sequence (e.g. 202600001).\n\n"
        "[Challenge: Normal Distribution]\n"
        "score = 80 + 10 * random.nextGaussian()\n"
        "limitScore() clamps result to [0, 100].\n\n"
        "[Key Code]\n"
        "public String generateId() {\n"
        "    int maxId = 0;\n"
        "    for (Student s : studentList) {\n"
        "        int num = Integer.parseInt(s.getId().substring(4));\n"
        "        if (num > maxId) maxId = num;\n"
        "    }\n"
        "    return String.format(\"2026%05d\", maxId + 1);\n"
        "}"
    )

# 4.3 Data persistence
idx, p = find_para_containing("*数据持久化")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "3. Data Persistence Module\n\n"
        "[Approach]\n"
        "Uses local text file student_data.txt (no database - pure Java basics).\n\n"
        "[File Format (CSV)]\n"
        "One student per line, 7 comma-separated fields:\n"
        "ID, Name, Gender, Birthday, Math, Java, PE\n\n"
        "[Write Flow - saveDataToFile()]\n"
        "1. Create BufferedWriter + FileWriter -> student_data.txt\n"
        "2. Iterate studentList, join 7 attributes with commas\n"
        "3. Write each student as one line\n"
        "4. Close stream\n"
        "5. Called after every add/update/delete operation\n\n"
        "[Read Flow - loadDataFromFile()]\n"
        "1. Check if file exists, skip if not\n"
        "2. Create BufferedReader + FileReader\n"
        "3. Read line by line, split by comma\n"
        "4. Create Student object, set attributes\n"
        "5. Add to studentList\n"
        "6. Called in constructor on startup"
    )

# 4.4 Database (file-based)
idx, p = find_para_containing("数据库设计")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "4. Data Storage Design\n\n"
        "[Note]\n"
        "This project uses text file storage (student_data.txt) instead of a database, "
        "in accordance with the freshman-level course requirements. Advantages:\n"
        "1. No database installation/configuration needed\n"
        "2. No JDBC driver import required\n"
        "3. Data file can be viewed/edited with Notepad\n"
        "4. Fully within freshman Java course scope\n\n"
        "[Upgrade Path]\n"
        "For future upgrade, MySQL + JDBC or SQLite could replace file storage "
        "by modifying saveDataToFile() and loadDataFromFile() methods in ScoreManager, "
        "without changing the GUI layer."
    )

idx, p = find_para_containing("数据库操作代码")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "5. Database Operations Code\n\n"
        "Not applicable (file-based storage used).\n"
        "See saveDataToFile() and loadDataFromFile() methods in ScoreManager.java "
        "for data access implementation."
    )

# ============================================================
# Part 5: Summary
# ============================================================
print("Filling Part 5: Summary...")

idx, p = find_para_containing("遇到的问题")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "This project successfully implemented a complete student score management "
        "system with GUI, CRUD, data persistence, statistics, and report export.\n\n"
        "Problems encountered and solutions:\n\n"
        "1. Unique Student ID\n"
        "   Problem: How to guarantee no duplicate IDs?\n"
        "   Solution: Traverse all existing IDs, find max sequence + 1.\n\n"
        "2. Data Persistence\n"
        "   Problem: Data lost when program closes.\n"
        "   Solution: Auto-save after each operation; load on startup.\n\n"
        "3. Table-Data Sync\n"
        "   Problem: Table not updating after operations.\n"
        "   Solution: refreshTable() called after all CRUD/sort operations.\n\n"
        "4. Normal Distribution\n"
        "   Problem: Generated scores may exceed 0-100 range.\n"
        "   Solution: limitScore() clamps values to valid range.\n\n"
        "5. Anonymous Inner Class Compatibility\n"
        "   Problem: Lambda expressions incompatible with older JDK.\n"
        "   Solution: All event listeners use anonymous inner class syntax."
    )

# ============================================================
# Part 6: Future Work
# ============================================================
print("Filling Part 6: Future Work...")

idx, p = find_para_containing("不足之处")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "Current limitations and improvement plans:\n\n"
        "1. Limitations\n"
        "   - Text file storage, no concurrent user support\n"
        "   - Hardcoded login credentials\n"
        "   - No input validation (ID format, date format)\n"
        "   - Excel export is stub only (needs POI library)\n"
        "   - No data visualization charts\n"
        "   - No unit tests\n\n"
        "2. Improvement Plans\n"
        "   - Learn MySQL and migrate to relational database\n"
        "   - Learn JDBC/MyBatis for standard data access layer (DAO)\n"
        "   - Learn JFreeChart for visualization\n"
        "   - Learn JUnit for unit testing\n"
        "   - Learn Maven/Gradle for dependency management\n"
        "   - Learn Git branching for version control best practices"
    )

idx, p = find_para_containing("选修")
if idx is not None:
    clear_para(idx)
    paragraphs[idx].add_run(
        "Learning Plan:\n"
        "Next steps: study JDBC basics (Connection, Statement, ResultSet); "
        "learn MySQL CRUD; learn JUnit for testing; learn Git branching and "
        "team collaboration; review project code and refactor repeated "
        "statistics logic into shared methods."
    )

# ============================================================
# Fill the technology summary table
# ============================================================
print("Filling technology table...")
if doc.tables:
    table = doc.tables[0]
    new_rows = [
        ("1", "Java Swing GUI",
         "JFrame, JPanel, JTable, JButton, JTextField components.\n"
         "BorderLayout + GridLayout + FlowLayout.\n"
         "Anonymous inner class ActionListener."),
        ("2", "Object-Oriented Programming",
         "Student entity class (private fields + public getters/setters).\n"
         "Constructor overloading.\n"
         "MVC: Model -> Controller -> View.\n"
         "toString() override."),
        ("3", "Java I/O Persistence",
         "BufferedReader/FileReader for loading.\n"
         "BufferedWriter/FileWriter for saving.\n"
         "CSV format (String.split/comma join).\n"
         "Auto-save after operations, auto-load on startup."),
        ("4", "Collections & Algorithms",
         "ArrayList<Student> dynamic storage.\n"
         "Collections.sort() + Comparator for descending sort.\n"
         "For-each iteration for statistics.\n"
         "String.contains() fuzzy search."),
        ("5", "Normal Distribution Data",
         "Random.nextGaussian() standard normal distribution.\n"
         "Formula: score = 80 + 10 * nextGaussian().\n"
         "limitScore() clamps to [0, 100].\n"
         "One-click 100K record generation with auto-save."),
    ]

    for row_idx, (num, feature, desc) in enumerate(new_rows):
        if row_idx + 1 < len(table.rows):
            target_row = table.rows[row_idx + 1]
            for ci, val in enumerate([num, feature, desc]):
                if ci < len(target_row.cells):
                    cell = target_row.cells[ci]
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ""
                    if cell.paragraphs:
                        cell.paragraphs[0].add_run(val)

# ============================================================
# Save
# ============================================================
doc.save(OUTPUT)
print(f"\nReport saved to: {OUTPUT}")
